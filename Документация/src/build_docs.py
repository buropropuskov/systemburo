#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Сборка комплекта документации из Markdown в .docx и .pdf.

Источник правды - .md рядом с готовыми файлами. Оформление: Times New Roman 14,
полуторный интервал, поля 30/15/20/20 мм.

Подписи нумеруются автоматически по номеру раздела: «Таблица 5.1 - Название»,
«Рисунок 4.2 - Название», «Листинг 6.1 - Название». В исходном тексте номер не
указывается, только название, поэтому вставка нового объекта не ломает нумерацию.

Оглавление собирается в два прохода: документ верстается без номеров страниц,
конвертируется в PDF, номера вычитываются из его текста, документ пересобирается.
Поле оглавления Word не используется: при конвертации без открытия в редакторе
оно осталось бы пустым.

Запуск:
    python3 Документация/src/build_docs.py            # весь комплект
    python3 Документация/src/build_docs.py deploy     # один документ
"""
import os
import re
import subprocess
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SYSTEM = "СИСТЕМА ЭЛЕКТРОННОЙ ПОДАЧИ ЗАЯВОК «БЮРО ПРОПУСКОВ»"
VERSION = "1.0"
DOC_DATE = "28.07.2026"

DOCS = {
    "overview": {
        "file": "Техническое описание системы",
        "title": "ТЕХНИЧЕСКОЕ ОПИСАНИЕ СИСТЕМЫ",
        "subtitle": "Назначение, возможности, устройство и защита информации",
        "audience": "Аудитория: руководители подразделений, специалисты по информационным "
                    "технологиям, специалисты по защите информации",
    },
    "deploy": {
        "file": "Руководство по развёртыванию и сопровождению",
        "title": "РУКОВОДСТВО ПО РАЗВЁРТЫВАНИЮ И СОПРОВОЖДЕНИЮ",
        "subtitle": "Установка, настройка и эксплуатация",
        "audience": "Аудитория: специалисты подразделения информационных технологий, "
                    "системные администраторы",
    },
    "pilot": {
        "file": "Критерии завершения пилотного периода",
        "title": "КРИТЕРИИ УСПЕШНОГО ЗАВЕРШЕНИЯ ПИЛОТНОГО ПЕРИОДА",
        "subtitle": "Условия проведения, проверяемые сценарии и порядок приёмки",
        "audience": "Аудитория: руководители подразделений, бюро пропусков, "
                    "подразделение информационных технологий",
    },
}

FONT = "Times New Roman"
MONO = "Courier New"
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)
CODE_FILL = "F4F4F4"
CODE_BORDER = "C8C8C8"
HEAD_FILL = "E8E8E8"

# Пастельные заливки блоков внимания. Насыщенность намеренно низкая: документ
# печатают на монохромных принтерах, и яркий фон превратился бы в серую плашку.
CALLOUTS = {
    "ОПАСНО": ("FBEAE7", "C0392B"),
    "ВАЖНО": ("FDF6E3", "B7791F"),
    "ПРИМЕЧАНИЕ": ("EAF1F8", "2C6FA6"),
}

TEXT_WIDTH_CM = 16.5
HEADING_SIZE = {1: 16, 2: 14, 3: 14, 4: 14}
CHAR_CM = 0.25
MONO_CHAR_CM = 0.26
CELL_PADDING_CM = 0.5


# --------------------------------------------------------------------------
# низкоуровневые помощники
# --------------------------------------------------------------------------
def set_font(run, name=FONT, size=14, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def shade_paragraph(paragraph, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def border_paragraph(paragraph, color=CODE_BORDER, sides=("top", "left", "bottom", "right"), size="6"):
    pbdr = OxmlElement("w:pBdr")
    for edge in sides:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        pbdr.append(el)
    paragraph.paragraph_format.element.get_or_add_pPr().append(pbdr)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    run._r.append(fld)

    run = paragraph.add_run("1")
    set_font(run, size=11, color=GREY)

    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)


def header_row_repeat(row):
    """Повтор шапки на каждой странице и запрет её отрыва от первой строки данных."""
    tr_pr = row._tr.get_or_add_trPr()
    for tag in ("w:tblHeader", "w:cantSplit"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "true")
        tr_pr.append(el)
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = True


def no_split_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def apply_grid(table, widths_cm):
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for col in list(grid):
        grid.remove(col)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 567))))
        grid.append(col)


# --------------------------------------------------------------------------
# оформление
# --------------------------------------------------------------------------
def page_setup(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)


def init_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(14)
    normal.font.color.rgb = BLACK
    rfonts = normal.element.rPr.rFonts
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Mm(12.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level in (1, 2, 3, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(HEADING_SIZE[level])
        style.font.bold = True
        style.font.color.rgb = BLACK
        rf = style.element.rPr.rFonts
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), FONT)
        hpf = style.paragraph_format
        hpf.first_line_indent = Mm(0)
        hpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hpf.line_spacing = 1.5
        hpf.space_before = Pt(14 if level == 1 else 10)
        hpf.space_after = Pt(6)
        hpf.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(14)
        style.font.color.rgb = BLACK
        rf = style.element.rPr.rFonts
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), FONT)
        lpf = style.paragraph_format
        lpf.line_spacing = 1.5
        lpf.first_line_indent = Mm(0)
        lpf.space_after = Pt(0)
        lpf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def plain(doc, text, size=14, bold=False, italic=False, align=None, color=BLACK):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Mm(0)
    if align is not None:
        para.alignment = align
    if text:
        set_font(para.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return para


def header_footer(doc):
    section = doc.sections[-1]
    section.different_first_page_header_footer = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Mm(0)
    add_field(footer, "PAGE")


def cover(doc, cfg):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Mm(0)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    set_font(para.add_run(SYSTEM), size=13, bold=True)

    for _ in range(6):
        plain(doc, "")
    plain(doc, cfg["title"], size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    plain(doc, "")
    plain(doc, cfg["subtitle"], size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        plain(doc, "")
    plain(doc, cfg["audience"], size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    plain(doc, "")
    plain(doc, f"Версия документа: {VERSION}", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    plain(doc, f"Дата: {DOC_DATE}", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


# --------------------------------------------------------------------------
# разбор Markdown
# --------------------------------------------------------------------------
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*\n]+\*)")


def add_inline(paragraph, text, size=14, color=BLACK):
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`") and token.endswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name=MONO, size=size - 2, color=color)
        elif token.startswith("[") and "](" in token:
            label, url = token[1:-1].split("](", 1)
            set_font(paragraph.add_run(label), size=size, color=color)
            if not url.startswith("#"):
                set_font(paragraph.add_run(f" ({url})"), size=size - 2, color=GREY)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True, color=color)
        else:
            set_font(paragraph.add_run(token), size=size, color=color)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def longest_word_cm(text):
    widest = 0.0
    for chunk in re.split(r"\s+", text.strip()):
        mono = chunk.startswith("`") or chunk.endswith("`")
        clean = chunk.strip("`*[]()")
        widest = max(widest, len(clean) * (MONO_CHAR_CM if mono else CHAR_CM))
    return widest + CELL_PADDING_CM


def column_widths(headers, rows):
    count = len(headers)
    weights, minimums = [], []
    for i in range(count):
        cells = [headers[i]] + [r[i] if i < len(r) else "" for r in rows]
        longest = max((len(c) for c in cells), default=1)
        typical = sum(len(c) for c in cells) / max(len(cells), 1)
        weights.append(max(longest * 0.35 + typical * 0.65, 4))
        minimums.append(max(longest_word_cm(c) for c in cells))

    if sum(minimums) > TEXT_WIDTH_CM:
        squeeze = TEXT_WIDTH_CM / sum(minimums)
        return [m * squeeze for m in minimums]

    widths = [0.0] * count
    free = set(range(count))
    remaining = TEXT_WIDTH_CM
    while True:
        total = sum(weights[i] for i in free)
        if total <= 0:
            break
        under = [i for i in free if weights[i] / total * remaining < minimums[i]]
        if not under:
            for i in free:
                widths[i] = weights[i] / total * remaining
            break
        for i in under:
            widths[i] = minimums[i]
            remaining -= minimums[i]
            free.discard(i)
        if not free:
            break
    scale = TEXT_WIDTH_CM / sum(widths)
    return [w * scale for w in widths]


def add_caption(doc, text, above=True):
    """Подпись к таблице (сверху) либо к рисунку (снизу).

    По ГОСТ подпись набирается курсивом, номер отделяется от названия длинным тире.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.first_line_indent = Mm(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(8 if above else 2)
    pf.space_after = Pt(4 if above else 10)
    pf.keep_with_next = above
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT if above else WD_ALIGN_PARAGRAPH.CENTER
    set_font(para.add_run(text), size=12, italic=True)



def add_list_item(doc, marker, text):
    """Пункт списка с явным маркером.

    Встроенные стили списков ведут сквозную нумерацию через весь документ, из-за
    чего второй список начинался бы не с единицы. Поэтому номер ставится вручную.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.first_line_indent = Mm(-8)
    pf.left_indent = Mm(20)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(para.add_run(marker + "\t"), size=14)
    add_inline(para, text)
    return para



def clean_cell(text):
    """Текст ячейки без разметки: в PDF попадает именно он."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()



def cell_probe(text):
    """Кусок ячейки, который заведомо лежит на одной строке колонки.

    Ячейка переносится внутри своей колонки, и в текстовом слое PDF между её
    частями оказывается текст соседних колонок. Поэтому берём первое слово, а
    если оно слишком короткое для поиска - добавляем следующее.
    """
    words = clean_cell(text).split()
    if not words:
        return ""
    probe = words[0]
    for word in words[1:]:
        if len(probe) >= 6:
            break
        probe = f"{probe} {word}"
    return probe


def split_rows(rows, breaks):
    """Разрезать строки таблицы в местах, где страница её разрывает."""
    if not breaks:
        return [rows]
    chunks, start = [], 0
    for point in breaks:
        if 0 < point < len(rows):
            chunks.append(rows[start:point])
            start = point
    chunks.append(rows[start:])
    return [c for c in chunks if c]


def find_first_unsplit_break(cfg, tables_meta, handled):
    """Найти первую сверху таблицу, которую страница рвёт, и место разрыва.

    Обрабатываем строго по одной сверху вниз. Вёрстка однонаправленна: содержимое
    ниже не влияет на положение того, что выше, поэтому уже зафиксированное
    разбиение от последующих не сдвинется. Попытка разобрать все таблицы за один
    проход не сходится: разбиение верхней таблицы смещает все нижние.

    Уместившиеся таблицы намеренно не запоминаются: разбиение вышележащей сдвигает
    вёрстку, и таблица, которая помещалась, начинает рваться.
    """
    pages = subprocess.run(
        ["pdftotext", "-layout", doc_path(cfg, ".pdf"), "-"],
        capture_output=True, text=True, check=True,
    ).stdout.split("\f")
    flat = [re.sub(r"\s+", " ", pg) for pg in pages]

    for meta in tables_meta:
        number = meta["number"]
        if number in handled:
            continue
        cells = [cell_probe(c) for c in meta["first_cells"]]
        if len(cells) < 3:
            continue

        marker = f"Таблица {number} "
        start = next((i for i, pg in enumerate(flat) if marker in pg), None)
        if start is None:
            continue

        pos, cursor, page = 0, flat[start].find(marker), start
        while page < len(flat) and pos < len(cells):
            matched = 0
            while pos < len(cells):
                cell = cells[pos]
                if not cell:
                    pos += 1
                    matched += 1
                    continue
                at = flat[page].find(cell, cursor)
                if at < 0:
                    break
                cursor, pos, matched = at + len(cell), pos + 1, matched + 1
            if pos >= len(cells):
                break
            if matched == 0:
                break
            # таблица не поместилась: строка pos начинается уже на следующей странице
            return number, [pos]
    return None, None



def continuation_at_page_top(cfg, number):
    """Проверить, что подпись «Продолжение таблицы» стоит в начале страницы.

    После разреза подпись занимает место, и остаток таблицы иногда перестаёт
    выходить за лист. Разрыв исчезает, а подпись повисает в середине страницы -
    такое разбиение надо откатить.
    """
    pages = subprocess.run(
        ["pdftotext", "-layout", doc_path(cfg, ".pdf"), "-"],
        capture_output=True, text=True, check=True,
    ).stdout.split("\f")
    needle = f"Продолжение таблицы {number}"
    for page in pages:
        lines = [ln for ln in page.splitlines() if ln.strip()]
        for idx, line in enumerate(lines):
            if needle in line:
                return idx == 0
    return False


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fixed_layout(table)
    widths = column_widths(headers, rows)

    head_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        shade_cell(head_cells[i], HEAD_FILL)
        para = head_cells[i].paragraphs[0]
        para.paragraph_format.first_line_indent = Mm(0)
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(para, title, size=12)
        for run in para.runs:
            run.font.bold = True
    header_row_repeat(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        no_split_row(table.rows[-1])
        for i in range(len(headers)):
            para = cells[i].paragraphs[0]
            para.paragraph_format.first_line_indent = Mm(0)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(para, row[i] if i < len(row) else "", size=12)

    apply_grid(table, widths)
    for r in table.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Mm(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 1.0
    set_font(spacer.add_run(""), size=6)


def add_code_block(doc, lines):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.first_line_indent = Mm(0)
    pf.left_indent = Mm(4)
    pf.right_indent = Mm(2)
    pf.line_spacing = 1.0
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.keep_together = True
    shade_paragraph(para, CODE_FILL)
    border_paragraph(para)
    for i, line in enumerate(lines):
        if i:
            para.add_run().add_break()
        set_font(para.add_run(line), name=MONO, size=10)


def add_callout(doc, kind, lines):
    """Блок повышенного внимания: заливка, толстая полоса слева, подпись типа."""
    fill, accent = CALLOUTS[kind]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    fixed_layout(table)
    apply_grid(table, [TEXT_WIDTH_CM])
    cell = table.cell(0, 0)
    cell.width = Cm(TEXT_WIDTH_CM)
    shade_cell(cell, fill)

    borders = OxmlElement("w:tcBorders")
    for edge, width, color in (
        ("left", "24", accent), ("top", "4", fill),
        ("bottom", "4", fill), ("right", "4", fill),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), width)
        el.set(qn("w:color"), color)
        borders.append(el)
    cell._tc.get_or_add_tcPr().append(borders)

    # Блок читается как единое предупреждение, разрыв между страницами его ломает.
    no_split_row(table.rows[0])

    first = cell.paragraphs[0]
    first.paragraph_format.first_line_indent = Mm(0)
    first.paragraph_format.line_spacing = 1.2
    first.paragraph_format.space_after = Pt(0)
    set_font(first.add_run(kind + ". "), size=13, bold=True,
             color=RGBColor.from_string(accent))
    # После точки в маркере идёт новое предложение, поэтому первая буква заглавная
    # независимо от того, как написано в исходном тексте.
    head = lines[0]
    if head and head[0].islower():
        head = head[0].upper() + head[1:]
    add_inline(first, head, size=13)

    for extra in lines[1:]:
        para = cell.add_paragraph()
        para.paragraph_format.first_line_indent = Mm(0)
        para.paragraph_format.line_spacing = 1.2
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(0)
        add_inline(para, extra, size=13)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Mm(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 1.0
    set_font(spacer.add_run(""), size=6)


def add_image(doc, path, caption_text):
    if not os.path.exists(path):
        raise FileNotFoundError(f"не найдено изображение: {path}")
    doc.add_picture(path, width=Cm(TEXT_WIDTH_CM))
    pic = doc.paragraphs[-1]
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.first_line_indent = Mm(0)
    pic.paragraph_format.space_before = Pt(6)
    pic.paragraph_format.space_after = Pt(2)
    pic.paragraph_format.keep_with_next = True
    add_caption(doc, caption_text, above=False)


def parse_markdown(md_text):
    blocks = []
    lines = md_text.split("\n")
    i = 0
    pending_caption = None

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped == "<!-- PAGEBREAK -->":
            blocks.append(("pagebreak", None))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # подпись к следующему объекту: «Таблица: ...» либо «Листинг: ...»
        cap = re.match(r"^(Таблица|Листинг):\s+(.+)$", stripped)
        if cap:
            pending_caption = (cap.group(1), cap.group(2).strip())
            i += 1
            continue

        # блок внимания: «> ВАЖНО: текст», продолжение строками «> ...»
        callout = re.match(r"^>\s*(ОПАСНО|ВАЖНО|ПРИМЕЧАНИЕ):\s*(.*)$", stripped)
        if callout:
            kind = callout.group(1)
            body = [callout.group(2).strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip("> ").strip())
                i += 1
            blocks.append(("callout", (kind, [b for b in body if b])))
            continue

        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", (code, pending_caption)))
            pending_caption = None
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            blocks.append(("heading", (len(heading.group(1)), heading.group(2).strip())))
            i += 1
            continue

        image = re.match(r"^!\[(.*?)\]\((.+?)\)$", stripped)
        if image:
            blocks.append(("image", (image.group(2), image.group(1))))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1].strip()
        ):
            headers = split_table_row(lines[i])
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append(("table", (headers, rows, pending_caption)))
            pending_caption = None
            continue

        bullet = re.match(r"^([-*])\s+(.*)$", stripped)
        if bullet:
            items = []
            while i < len(lines):
                m = re.match(r"^([-*])\s+(.*)$", lines[i].strip())
                if not m or not lines[i].strip():
                    break
                text = m.group(2)
                i += 1
                while i < len(lines) and lines[i].startswith("  ") and not re.match(
                    r"^\s*([-*]|\d+\.)\s", lines[i]
                ):
                    text += " " + lines[i].strip()
                    i += 1
                items.append(text)
            blocks.append(("bullets", items))
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            items = []
            while i < len(lines):
                m = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if not m or not lines[i].strip():
                    break
                text = m.group(1)
                i += 1
                while i < len(lines) and lines[i].startswith("   ") and not re.match(
                    r"^\s*(\d+\.|[-*])\s", lines[i]
                ):
                    text += " " + lines[i].strip()
                    i += 1
                items.append(text)
            blocks.append(("numbers", items))
            continue

        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,4}\s|```|\||[-*]\s|\d+\.\s|!\[|>\s*(ОПАСНО|ВАЖНО|ПРИМЕЧАНИЕ):|(Таблица|Листинг):\s)",
            lines[i].strip(),
        ):
            para.append(lines[i].strip())
            i += 1
        blocks.append(("para", " ".join(para)))

    return blocks


def section_prefix(title):
    """Префикс нумерации объектов: «5» из «5. Требования», «А» из «Приложение А»."""
    app = re.match(r"^Приложение\s+([А-ЯA-Z])", title)
    if app:
        return app.group(1)
    num = re.match(r"^(\d+)\.", title)
    return num.group(1) if num else None


def render(doc, blocks, toc_titles, splitmap=None, tables_meta=None):
    counters = {"table": 0, "figure": 0, "listing": 0}
    prefix = "1"
    splitmap = splitmap or {}

    for kind, payload in blocks:
        if kind == "pagebreak":
            doc.add_page_break()

        elif kind == "heading":
            level, text = payload
            if level == 1:
                new_prefix = section_prefix(text)
                if new_prefix:
                    prefix = new_prefix
                    counters = {"table": 0, "figure": 0, "listing": 0}
            para = doc.add_heading(level=min(level, 4))
            para.paragraph_format.first_line_indent = Mm(0)
            set_font(para.add_run(text), size=HEADING_SIZE[min(level, 4)], bold=True)
            if level <= 2:
                toc_titles.append((level, text))

        elif kind == "para":
            add_inline(doc.add_paragraph(), payload)

        elif kind == "bullets":
            for item in payload:
                add_list_item(doc, "\u2022", item)

        elif kind == "numbers":
            for pos, item in enumerate(payload, 1):
                add_list_item(doc, f"{pos}.", item)

        elif kind == "callout":
            add_callout(doc, payload[0], payload[1])

        elif kind == "code":
            code, cap = payload
            if cap:
                counters["listing"] += 1
                add_caption(doc, f"Листинг {prefix}.{counters['listing']} \u2014 {cap[1]}")
            add_code_block(doc, code)

        elif kind == "table":
            headers, rows, cap = payload
            if cap:
                counters["table"] += 1
                number = f"{prefix}.{counters['table']}"
                if tables_meta is not None:
                    tables_meta.append({"number": number,
                                        "first_cells": [r[0] if r else "" for r in rows]})
                for idx, chunk in enumerate(split_rows(rows, splitmap.get(number))):
                    if idx == 0:
                        add_caption(doc, f"Таблица {number} \u2014 {cap[1]}")
                    else:
                        # Принудительный разрыв не нужен: часть выше закончилась на
                        # границе листа, а подпись держится за своей таблицей.
                        add_caption(doc, f"Продолжение таблицы {number}")
                    add_table(doc, headers, chunk)
            else:
                add_table(doc, headers, rows)

        elif kind == "image":
            path, caption = payload
            if not os.path.isabs(path):
                path = os.path.join(BASE_DIR, path)
            counters["figure"] += 1
            add_image(doc, path, f"Рисунок {prefix}.{counters['figure']} \u2014 {caption}")


def render_toc(doc, entries, pagemap):
    heading = doc.add_heading(level=1)
    heading.paragraph_format.first_line_indent = Mm(0)
    set_font(heading.add_run("СОДЕРЖАНИЕ"), size=16, bold=True)

    for level, title in entries:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.first_line_indent = Mm(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.5
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if level == 2:
            pf.left_indent = Mm(8)
        pf.tab_stops.add_tab_stop(Cm(TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        set_font(para.add_run(title), size=13, bold=(level == 1))
        page = pagemap.get(title, "") if pagemap else ""
        set_font(para.add_run("\t" + str(page)), size=13)
    doc.add_page_break()


def doc_dir(cfg):
    """Каждый документ лежит в собственной папке: .md, .docx и .pdf вместе."""
    path = os.path.join(BASE_DIR, cfg["file"])
    os.makedirs(path, exist_ok=True)
    return path


def doc_path(cfg, ext):
    return os.path.join(doc_dir(cfg), cfg["file"] + ext)


def build(cfg, pagemap=None, splitmap=None):
    src = doc_path(cfg, ".md")
    with open(src, encoding="utf-8") as fh:
        blocks = parse_markdown(fh.read())

    probe = Document()
    page_setup(probe.sections[0])
    init_styles(probe)
    toc_titles = []
    render(probe, blocks, toc_titles)

    doc = Document()
    page_setup(doc.sections[0])
    init_styles(doc)
    header_footer(doc)
    cover(doc, cfg)
    render_toc(doc, toc_titles, pagemap)
    tables_meta = []
    render(doc, blocks, [], splitmap, tables_meta)
    doc.save(doc_path(cfg, ".docx"))
    return toc_titles, tables_meta


def to_pdf(cfg):
    docx = doc_path(cfg, ".docx")
    pdf = doc_path(cfg, ".pdf")
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", doc_dir(cfg), docx],
        capture_output=True, text=True, timeout=300,
    )
    if not os.path.exists(pdf):
        raise RuntimeError(f"конвертация в PDF не удалась: {result.stdout} {result.stderr}")
    return pdf


def pagemap_from_pdf(cfg, entries):
    pdf = doc_path(cfg, ".pdf")
    text = subprocess.run(
        ["pdftotext", "-layout", pdf, "-"], capture_output=True, text=True, check=True
    ).stdout
    pages = text.split("\f")

    # Оглавление занимает несколько страниц и содержит те же заголовки. Строки с
    # точечным лидером отбрасываем, иначе раздел получит номер страницы оглавления.
    def body_only(page):
        return re.sub(r"\s+", " ", " ".join(
            ln for ln in page.splitlines() if "...." not in ln))

    start = 0
    for idx, page in enumerate(pages):
        if "СОДЕРЖАНИЕ" in page:
            start = idx + 1
            break

    pagemap = {}
    for _, title in entries:
        needle = re.sub(r"\s+", " ", title).strip()
        for idx in range(start, len(pages)):
            if needle in body_only(pages[idx]):
                pagemap[title] = idx + 1
                break
    return pagemap




def report_silent_breaks(cfg):
    """Предупредить о таблицах, которые страница рвёт без подписи продолжения.

    Итоговая проверка результата: подбор разбиения опирается на распознавание
    таблиц в текстовом слое PDF и в редком случае может таблицу пропустить.
    """
    with open(doc_path(cfg, ".md"), encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    heads = set()
    for i, line in enumerate(lines):
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1].strip()
        ):
            cells = [re.sub(r"[`*]", "", c).strip() for c in line.strip().strip("|").split("|")]
            if len(cells) >= 2 and cells[0]:
                heads.add(re.sub(r"\s+", " ", " ".join(cells)))

    pages = subprocess.run(
        ["pdftotext", "-layout", doc_path(cfg, ".pdf"), "-"],
        capture_output=True, text=True, check=True,
    ).stdout.split("\f")

    silent = []
    for number, page in enumerate(pages, 1):
        rows = [ln for ln in page.splitlines() if ln.strip()]
        if not rows or "Продолжение таблицы" in rows[0]:
            continue
        if re.sub(r"\s+", " ", rows[0]).strip() in heads:
            silent.append(number)
    if silent:
        print(f"  ВНИМАНИЕ: таблица разорвана без подписи продолжения, страницы: {silent}")
    return silent


def build_one(key):
    cfg = DOCS[key]
    name = cfg["file"]

    # Таблицы, которые страница разрывает, обрабатываются по одной сверху вниз:
    # фиксируем разбиение верхней, пересобираем, ищем следующую. Так каждое
    # разбиение определяется на вёрстке, где всё вышележащее уже окончательно.
    splitmap, handled = {}, set()
    entries, tables_meta = build(cfg, None, splitmap)
    to_pdf(cfg)

    for _ in range(60):
        number, breaks = find_first_unsplit_break(cfg, tables_meta, handled)
        if not number:
            break
        handled.add(number)

        # Подпись продолжения сама занимает строку, поэтому остаток иногда
        # перестаёт выходить за лист и подпись повисает в середине страницы.
        # В таком случае режем на строку раньше.
        placed = False
        for shift in (0, 1, 2):
            point = breaks[0] - shift
            if point < 1:
                break
            splitmap[number] = [point]
            entries, tables_meta = build(cfg, None, splitmap)
            to_pdf(cfg)
            if continuation_at_page_top(cfg, number):
                placed = True
                break
        if not placed:
            splitmap.pop(number, None)
            entries, tables_meta = build(cfg, None, splitmap)
            to_pdf(cfg)

    if splitmap:
        print(f"  таблицы с продолжением: {', '.join(sorted(splitmap))}")

    pagemap = pagemap_from_pdf(cfg, entries)
    missing = [t for _, t in entries if t not in pagemap]
    if missing:
        print(f"  ВНИМАНИЕ: не найдено в PDF {len(missing)}: {missing[:3]}")

    print(f"[{name}] финальная сборка...")
    build(cfg, pagemap, splitmap)
    pdf = to_pdf(cfg)
    report_silent_breaks(cfg)
    pages = subprocess.run(["pdfinfo", pdf], capture_output=True, text=True).stdout
    total = re.search(r"Pages:\s+(\d+)", pages)
    print(f"[{cfg['file']}] готово, страниц: {total.group(1) if total else '?'}")


def main():
    keys = sys.argv[1:] or list(DOCS)
    for key in keys:
        if key not in DOCS:
            sys.exit(f"неизвестный документ: {key}. Доступны: {', '.join(DOCS)}")
        build_one(key)


if __name__ == "__main__":
    main()
